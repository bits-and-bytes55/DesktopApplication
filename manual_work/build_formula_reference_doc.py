from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\DesktopApplication")
OUTPUT = ROOT / "MudPro-Desktop-App" / "docs" / "MudPro_Formula_Reference_v1.0.docx"
OUTPUT_GDOCS = (
    ROOT / "MudPro-Desktop-App" / "docs" / "MudPro_Formula_Reference_v1.0_gdocs.docx"
)


SECTIONS = [
    {
        "title": "1. Scope And Conventions",
        "rows": [
            (
                "Document scope",
                [
                    "Formula: This document covers the calculation logic currently implemented in the MudPro desktop app and backend services.",
                    "Notes: It focuses on engineering, report, pit-volume, operation, hydraulics, survey, concentration, and solids-analysis formulas. CRUD-only screens and pure unit-conversion tables are excluded.",
                    "Source: MudPro-Desktop-App/lib, Mud-pro-Backend/src",
                ],
            ),
            (
                "Rounding",
                [
                    "Formula: Most report-volume values use round2(value) = Number(value.toFixed(2)); some nozzle and product values use 3 or 4 decimals.",
                    "Notes: UI displays usually follow bbl = 2-3 decimals, nozzle diameter = 4 decimals, pump displacement = 4 decimals, pump rate = 1 decimal.",
                    "Source: volumeName.controller.js, nozzle.controller.js, pump.controller.js, consumeProductController.js",
                ],
            ),
            (
                "Base units",
                [
                    "Formula: Core formulas are evaluated in base oilfield units such as ft, in, bbl, gpm, psi, ppg, and lb/ft before AppUnits converts them for display.",
                    "Notes: This is why many UI screens first convert user input to a fixed base unit and convert the result back to the selected unit system.",
                    "Source: AppUnits usage across well_tab_content.dart, pump_tab_content.dart, formation_controller.dart, engineering_tools_controller.dart, bit_hydra_controller.dart",
                ],
            ),
        ],
    },
    {
        "title": "2. Well, Hole, And Pipe Geometry",
        "rows": [
            (
                "Raw cylinder volume",
                [
                    "Formula: Volume_bbl = (ID_in^2 x Length_ft) / 1029.4",
                    "Notes: This is the base cylinder formula used for pipe IDs, pipe ODs, casing IDs, and hole intervals.",
                    "Source: pitvolumename/volumeName.controller.js rawCylinderVolume(), Export/exportInventoryController.js calculatePipeVolume()",
                ],
            ),
            (
                "Cased-hole length",
                [
                    "Formula: Length_ft = abs(Shoe - Top) when both are valid; otherwise max(0, min(MD, Shoe) - Top) or max(0, MD - Top)",
                    "Notes: The code supports full casing rows and liner rows.",
                    "Source: pitvolumename/volumeName.controller.js casedHoleLength()",
                ],
            ),
            (
                "Cased-hole raw volume",
                [
                    "Formula: CasedHole_bbl = sum(rawCylinderVolume(ID, Length)) across valid casing rows",
                    "Notes: If liner rows exist, the controller also adds the previous casing volume up to the first liner top.",
                    "Source: pitvolumename/volumeName.controller.js calculateCasedHoleRawVolume()",
                ],
            ),
            (
                "Open-hole volume with washout",
                [
                    "Formula: EffectiveHoleID_in = HoleID_in x (1 + Washout_pct / 100)",
                    "Formula: OpenHole_bbl = sum(rawCylinderVolume(EffectiveHoleID, MD_i - PreviousDepth))",
                    "Notes: PreviousDepth starts at the deepest cased-hole shoe.",
                    "Source: pitvolumename/volumeName.controller.js calculateOpenHoleRawVolume()",
                ],
            ),
            (
                "Drill-string steel and inside volume",
                [
                    "Formula: PipeSteel_bbl = sum(rawCylinderVolume(OD, PieceLength))",
                    "Formula: PipeInside_bbl = sum(rawCylinderVolume(ID, PieceLength))",
                    "Notes: PieceLength is clipped to the depth limit, usually bit depth or MD.",
                    "Source: pitvolumename/volumeName.controller.js calculateDrillStringGuideVolumes()",
                ],
            ),
            (
                "Displacement and hole volume",
                [
                    "Formula: HoleSpace_bbl = CasedHole_bbl + OpenHole_bbl",
                    "Formula: Displacement_bbl = max(0, PipeSteel_bbl - PipeInside_bbl)",
                    "Formula: Hole_bbl = HoleSpace_bbl - PipeSteel_bbl + PipeInside_bbl",
                    "Notes: The report uses Hole as the visible fluid-holding capacity after string displacement.",
                    "Source: pitvolumename/volumeName.controller.js calculateCombinedHoleVolumeResult()",
                ],
            ),
            (
                "Annulus in pit snapshot",
                [
                    "Formula: Annulus_bbl = max(0, Hole_bbl - DrillStringInside_bbl)",
                    "Notes: DrillStringInside is the counted pipe internal volume from the hole-volume result.",
                    "Source: pitvolumename/volumeName.controller.js lines building pit snapshot payload",
                ],
            ),
            (
                "Top plug",
                [
                    "Formula: Capacity_bbl_per_ft = 0.0009714 x HoleDiameter_in^2",
                    "Formula: PlugTop_ft = max(MD_ft - (CementPlugVol_bbl / Capacity_bbl_per_ft), 0)",
                    "Notes: UI converts the entered volume and length into bbl and ft before calculating.",
                    "Source: dashboard/tabs/well_tab_content.dart _calculateTopPlug()",
                ],
            ),
            (
                "Drill-string ID from OD and weight",
                [
                    "Formula: ID_in = sqrt(OD_in^2 - (Weight_lb_per_ft / 2.672))",
                    "Notes: Used by the report Well tab helper when OD and linear weight are known.",
                    "Source: dashboard/tabs/well_tab_content.dart _calculateIds()",
                ],
            ),
        ],
    },
    {
        "title": "3. Nozzle And Bit Calculations",
        "rows": [
            (
                "Single nozzle geometry",
                [
                    "Formula: Diameter_in = Size_32nds / 32",
                    "Formula: Area_in2 = (pi x Diameter_in^2) / 4",
                    "Notes: Backend stores both diameter and area per processed row.",
                    "Source: controllers/nozzle/nozzle.controller.js calculateNozzleArea()",
                ],
            ),
            (
                "Total flow area (TFA)",
                [
                    "Formula: TFA_in2 = sum(RawArea_in2 x Count) across all valid nozzle rows",
                    "Notes: The frontend nozzle controller mirrors the backend formula locally.",
                    "Source: controllers/nozzle/nozzle.controller.js processNozzles(), dashboard/controller/nozzle_controller.dart recalculateTfa()",
                ],
            ),
            (
                "Bit hydraulics utility",
                [
                    "Formula: TotalJetArea_in2 = sum(0.785 x Dia_in^2)",
                    "Formula: NozzleVelocity_ft_per_s = (0.408 x Q_gpm) / TotalJetArea",
                    "Formula: BitPressureDrop_psi = 0.65 x StandpipePressure_psi",
                    "Formula: HydraulicHP = (BitPressureDrop_psi x Q_gpm) / 1714",
                    "Formula: BitArea_in2 = 0.785 x BitSize_in^2",
                    "Formula: HHP_per_Area = HydraulicHP / BitArea",
                    "Formula: PressureDrop_pct = (BitPressureDrop / StandpipePressure) x 100",
                    "Formula: JetImpact_lbf = 0.01823 x MW_ppg x Q_gpm x NozzleVelocity",
                    "Source: utility/controller/bit_hydra_controller.dart calculateBitHydraulics()",
                ],
            ),
        ],
    },
    {
        "title": "4. Pump, Circulation, And Hydraulics",
        "rows": [
            (
                "Pump displacement",
                [
                    "Formula: Duplex with rod = 0.000162 x (2 x D^2 - d^2) x L x Eff",
                    "Formula: Duplex without rod = 0.000324 x D^2 x L x Eff",
                    "Formula: Triplex = 0.000243 x D^2 x L x Eff",
                    "Formula: Quadplex = 0.000324 x D^2 x L x Eff",
                    "Formula: Quintuplex = 0.000405 x D^2 x L x Eff",
                    "Notes: Eff is efficiency expressed as a fraction, not percent.",
                    "Source: controllers/pump/pump.controller.js calculateDisplacement(), UG/model/pump_model.dart recalculateDisplacement()",
                ],
            ),
            (
                "Pump rate",
                [
                    "Formula: Rate_gpm = Displacement_bbl_per_stk x SPM x 42",
                    "Notes: Both backend and report UI use the same oilfield conversion factor.",
                    "Source: controllers/pump/pump.controller.js calculateRate(), dashboard/tabs/pump_tab_content.dart recalculateRate()",
                ],
            ),
            (
                "Circulation timing",
                [
                    "Formula: Strokes = Volume_bbl / Displacement_bbl_per_stk",
                    "Formula: Minutes = (Volume_bbl x 42) / PumpRate_gpm",
                    "Fallback: Minutes = Strokes / StrokePumpSPM when pump rate is unavailable",
                    "Source: Export/exportInventoryController.js calculateCirculationTiming()",
                ],
            ),
            (
                "Annular velocity utility",
                [
                    "Formula: AV_ft_per_min = (24.51 x Q_gpm) / (Dh_in^2 - Dp_in^2)",
                    "Notes: The calculator first converts flow to gpm and diameters to inches.",
                    "Source: utility/controller/engineering_tools_controller.dart calculateAnnularVelocity()",
                ],
            ),
            (
                "Hydraulic ECD",
                [
                    "Formula: ECD_ppg = MW_ppg + (PressureLoss_psi / (0.052 x Depth_ft))",
                    "Notes: If pressure loss or depth is missing, the result falls back to base mud weight.",
                    "Source: Export/exportInventoryController.js calculateEcd(), report/recap hydra controller _calculateEcd()",
                ],
            ),
            (
                "Critical velocity",
                [
                    "Formula: Gap_in = HoleSize_in - PipeOD_in",
                    "Formula: CriticalVelocity = (1.08 x PV + 1.08 x sqrt(PV^2 + 12.34 x Gap^2 x YP x MW)) / (MW x Gap)",
                    "Source: Export/exportInventoryController.js hydraulicCriticalVelocity()",
                ],
            ),
            (
                "Pipe and annular velocity",
                [
                    "Formula: PipeVelocity_ft_per_min = (24.51 x Q_gpm) / ID_in^2",
                    "Formula: DrillStringVelocity_ft_per_s = PipeVelocity_ft_per_min / 60",
                    "Formula: AnnularVelocity_ft_per_min = (24.51 x Q_gpm) / (Dh_in^2 - OD_in^2)",
                    "Source: Export/exportInventoryController.js hydraulicPipeVelocity(), hydraulicDrillStringVelocity(), hydraulicAnnularVelocity()",
                ],
            ),
            (
                "Annular and drill-string pressure loss",
                [
                    "Formula: AnnularLoss = (PV x L x AV) / (60000 x Gap^2) + (YP x L) / (225 x Gap)",
                    "Formula: DrillStringLoss = (PV x L x PipeVel) / (PipeDenominator(ID) x ID^2) + (YP x L) / (225 x ID)",
                    "Notes: PipeDenominator(ID) is a polynomial derived from the legacy hydraulics model.",
                    "Source: Export/exportInventoryController.js hydraulicAnnularPressureLoss(), hydraulicDrillStringPressureWeight(), hydraulicPipePressureDenominator()",
                ],
            ),
            (
                "Bit pressure loss and jet velocity",
                [
                    "Formula: BitLoss_psi = (MW_ppg x Q_gpm^2) / (10858 x TFA_in2^2)",
                    "Formula: JetVelocity_ft_per_s = (0.32086 x Q_gpm) / TFA_in2",
                    "Source: Export/exportInventoryController.js hydraulicBitPressureLoss(), hydraulicBitJetVelocity()",
                ],
            ),
        ],
    },
    {
        "title": "5. Mud, Rheology, Solids, And Formation Windows",
        "rows": [
            (
                "Bingham rheology",
                [
                    "Formula: PV_cP = R600 - R300",
                    "Formula: YP_lbf_per_100ft2 = R300 - PV",
                    "Source: dashboard/controller/mud_controller.dart calculateRheology()",
                ],
            ),
            (
                "Power Law rheology",
                [
                    "Formula: n = 3.32 x log10(R600 / R300)",
                    "Formula: K = R600 / (1022^n)",
                    "Source: dashboard/controller/mud_controller.dart calculateRheology()",
                ],
            ),
            (
                "Herschel-Bulkley rheology",
                [
                    "Formula: YieldStress = max(0, 2 x R3 - R6)",
                    "Formula: Adjusted600 = R600 - YieldStress",
                    "Formula: Adjusted300 = R300 - YieldStress",
                    "Formula: n = 3.32 x log10(Adjusted600 / Adjusted300)",
                    "Formula: K = Adjusted600 / (1022^n)",
                    "Source: dashboard/controller/mud_controller.dart calculateRheology()",
                ],
            ),
            (
                "Formation window conversions",
                [
                    "Formula: Gradient_psi_per_ft = PPG x 0.052",
                    "Formula: Pressure_psi = Gradient x TVD_ft",
                    "Formula: PPG = Gradient / 0.052",
                    "Formula: Gradient = Pressure / TVD_ft",
                    "Source: UG/controller/formation_controller.dart _recalculateGroup()",
                ],
            ),
            (
                "Solids analysis - brine and corrected solids",
                [
                    "Formula: BrineSG = 0.99707 + 0.007923 x CaCl2_pct + 0.00004964 x CaCl2_pct^2",
                    "Formula: BrineVol_pct = (100 x Water_pct) / (BrineSG x (100 - CaCl2_pct) x 0.99707) when salt correction is needed",
                    "Formula: DissolvedSolids_pct = max(0, (BrineSG - 1) x 100)",
                    "Formula: CorrectedSolids_pct = 100 - (Oil_pct + BrineVol_pct) unless directly provided",
                    "Source: controllers/SolidAnalysis/solidanalysiscontroller.js computeSolidsAnalysis()",
                ],
            ),
            (
                "Solids analysis - average density and phase split",
                [
                    "Formula: AvgSolidsSG = (100 x (MW / 8.34) - Oil_pct x OilSG - BrineSG x BrineVol_pct) / CorrectedSolids_pct",
                    "Formula: HGS_pct = ((AvgSolidsSG - LGS_SG) / (HGS_SG - LGS_SG)) x CorrectedSolids_pct",
                    "Formula: LGS_pct = CorrectedSolids_pct - HGS_pct",
                    "Formula: LGS_lb_per_bbl = 3.5 x LGS_SG x LGS_pct",
                    "Formula: HGS_lb_per_bbl = 3.5 x HGS_SG x HGS_pct",
                    "Source: controllers/SolidAnalysis/solidanalysiscontroller.js computeSolidsAnalysis()",
                ],
            ),
            (
                "Solids analysis - bentonite and drill solids",
                [
                    "Formula: Bentonite_pct = Bentonite_lb_per_bbl / (3.5 x 2.65)",
                    "Formula: DrillSolids_pct = LGS_pct - Bentonite_pct",
                    "Formula: DrillSolids_lb_per_bbl = 3.5 x LGS_SG x DrillSolids_pct",
                    "Formula: DS_Bent_Ratio = DrillSolids_pct / Bentonite_pct",
                    "Formula: TotalSolids_lb_per_bbl = MW x 42 x (TotalSolids_pct / 100)",
                    "Source: controllers/SolidAnalysis/solidanalysiscontroller.js computeSolidsAnalysis()",
                ],
            ),
        ],
    },
    {
        "title": "6. Operations And Inventory Calculations",
        "rows": [
            (
                "Consume Product row math",
                [
                    "Formula: Final = Initial + Received - Returned - Adjust - Used",
                    "Formula: Cost = Used x Price",
                    "Formula: Volume_bbl is calculated only when the selected product is flagged for volume addition",
                    "Source: dashboard/tabs/operation/consume_product.dart ProductRowData.recalculate()",
                ],
            ),
            (
                "Consume Product unit-to-volume conversion",
                [
                    "Formula: gal -> bbl = TotalUnits / 42",
                    "Formula: bbl -> bbl = TotalUnits",
                    "Formula: kg -> bbl = TotalUnits / (SG x 158.987)",
                    "Formula: lb -> bbl = TotalUnits / (SG x 350)",
                    "Formula: ton or mt -> bbl = (TotalUnits x 2000) / (SG x 350)",
                    "Formula: liter -> bbl = TotalUnits / 158.987",
                    "Formula: ml -> bbl = TotalUnits / 158987",
                    "Formula: m3 -> bbl = TotalUnits x 6.28981",
                    "Source: controllers/Consumeproduct/consumeProductController.js calculateVolumeBbl(), dashboard/tabs/operation/consume_product.dart _calculateVolumeBbl()",
                ],
            ),
            (
                "Consume Product total volume",
                [
                    "Formula: TotalVol = sum(ProductCalculatedVolume) + WaterVolume_if_AddWaterChecked",
                    "Notes: This is the value shown in the Add Water panel and used for automatic Active System distribution.",
                    "Source: dashboard/tabs/operation/consume_product.dart _recalculateTotalVolume()",
                ],
            ),
            (
                "Distribution behavior",
                [
                    "Formula: If no manual distribution rows exist, DistributeRow[0].Volume = TotalVol",
                    "Notes: This preserves the legacy Active System mirroring behavior.",
                    "Source: dashboard/tabs/operation/consume_product.dart _rebalanceDistributeVolumes()",
                ],
            ),
            (
                "Receive Mud",
                [
                    "Formula: NetVolume = GrossVolume - LossVolume",
                    "Notes: MW, mud type, and leasing fee fall back to the selected premixed mud master when blank.",
                    "Source: controllers/receivemud/receiveMud.controller.js prepareReceiveMudData()",
                ],
            ),
            (
                "Return / Lost Mud",
                [
                    "Formula: TotalDeduct = VolReturned + VolLost",
                    "Notes: Source availability is validated before the transaction is accepted.",
                    "Source: controllers/returnlostmud/returnLostMud.controller.js prepareReturnLostMudData(), assertSourceHasAvailableVolume()",
                ],
            ),
            (
                "Other Volume Addition",
                [
                    "Formula: TotalVolume = Formation + Cuttings + VolumeNotFluid",
                    "Source: controllers/othervol/otherVolAddition.controller.js prepareOtherVolAdditionData()",
                ],
            ),
            (
                "Mud Loss - Active System",
                [
                    "Formula: TotalLoss = CuttingsRetention + Seepage + Dump + Shakers + Centrifuge + Evaporation + PitCleaning + Formation + AbandonInHole + LeftBehindCasing + Tripping + ExtraLoss",
                    "Notes: The backend validates that the Active System has enough available volume before accepting the loss.",
                    "Source: controllers/mudloss/mudLoss.controller.js createMudLoss()",
                ],
            ),
            (
                "Consume Service / Package / Engineering",
                [
                    "Formula: PackageFinal = Initial - Used",
                    "Formula: PackageCost = Used x Price",
                    "Formula: ServiceCost = Usage x Price",
                    "Formula: EngineeringCost = Usage x Price",
                    "Source: dashboard/tabs/operation/consume_service.dart",
                ],
            ),
            (
                "Cuttings/Retention helper dialog",
                [
                    "Formula: CuttingsRetention_bbl = VolDrilled_bbl x MudLossRatio_pct / 100",
                    "Source: dashboard/tabs/operation/operation_desktop_ui.dart showCuttingsRetentionDialog()",
                ],
            ),
            (
                "Evaporation helper dialog",
                [
                    "Formula: ActiveHours = DrillingHours + CirculatingHours",
                    "Formula: Evaporation_bbl = max(FlowlineTemp_F - 32, 0) x ActiveHours / 1000",
                    "Source: dashboard/tabs/operation/operation_desktop_ui.dart showEvaporationDialog()",
                ],
            ),
            (
                "Other Volume Addition helper dialog",
                [
                    "Formula: CuttingsGain = VolDrilled when ShakerBypass = Yes",
                    "Formula: CuttingsGain = VolDrilled x (1 - Efficiency_pct / 100) when ShakerBypass = No",
                    "Source: dashboard/tabs/operation/othervolumeaddition_view.dart",
                ],
            ),
        ],
    },
    {
        "title": "7. Pit, Volume Name, And Ledger Balancing",
        "rows": [
            (
                "Active Pits total",
                [
                    "Formula: ActivePits_bbl = sum(ActivePitMeasuredVolume)",
                    "Source: pitvolumename/volumeName.controller.js active pit reduction and snapshot builders",
                ],
            ),
            (
                "Active System",
                [
                    "Formula: ActiveSystem_bbl = Hole_bbl + ActivePitsWithTransfer_bbl",
                    "Notes: ActivePitsWithTransfer includes transfer deltas and consume-product distribution deltas by active pit.",
                    "Source: pitvolumename/volumeName.controller.js derivedActiveSystem",
                ],
            ),
            (
                "End volume",
                [
                    "Formula: EndVol_bbl = CarryOverBaseline + OperationLedgerDelta",
                    "Formula: OperationLedgerDelta = EndVolDelta + DistributedToActiveSystem + DistributedToActivePits",
                    "Exception: EndVol = 0 when the report has no current volume data and no carry-over, or when Empty Active System uses Dump.",
                    "Source: pitvolumename/volumeName.controller.js calculateEndVolForReport()",
                ],
            ),
            (
                "Pending difference",
                [
                    "Formula: Pending_bbl = EndVol_bbl - ActiveSystem_bbl",
                    "Notes: This is the 'End Vol. - Active System' value shown to the user for pit adjustment.",
                    "Source: pitvolumename/volumeName.controller.js endVolMinusActiveSystem",
                ],
            ),
            (
                "Total on location",
                [
                    "Formula: TotalOnLocation_bbl = ActiveSystem_bbl + AdjustedTotalStorage_bbl",
                    "Formula: AdjustedTotalStorage_bbl = TotalStorage_bbl + ReturnLostStorageDelta_bbl",
                    "Source: pitvolumename/volumeName.controller.js totalOnLocation",
                ],
            ),
            (
                "Ledger-style total on location",
                [
                    "Formula: LedgerTotal = ConsumeProductTotal + ReceivedMudTotal + AddWaterTotal + OtherVolAdditionTotal - LostMudTotal - MudLossTotal - MudLossStorageTotal",
                    "Notes: This is kept as a separate derived total in the pit snapshot payload.",
                    "Source: pitvolumename/volumeName.controller.js ledgerTotalOnLocation",
                ],
            ),
            (
                "Operation deltas that affect End Vol",
                [
                    "Formula: Add Water to Active System or active pit -> EndVolDelta += Volume",
                    "Formula: Receive Mud to Active System or active pit -> EndVolDelta += NetVolume",
                    "Formula: Return/Lost from Active System or active pit -> EndVolDelta -= (Returned + Lost)",
                    "Formula: Other Vol Addition -> EndVolDelta += TotalVolume",
                    "Formula: Mud Loss Active System -> EndVolDelta -= TotalLoss",
                    "Formula: Transfer from active to storage -> EndVolDelta -= TransferVolume",
                    "Formula: Transfer from storage to active -> EndVolDelta += TransferVolume",
                    "Formula: Empty Active System -> Transfer to Storage subtracts from EndVol; Dump forces EndVol to zero",
                    "Source: pitvolumename/volumeName.controller.js buildOperationVolumeEffects()",
                ],
            ),
        ],
    },
    {
        "title": "8. Survey, Concentration, And Cost Splits",
        "rows": [
            (
                "Survey minimum-curvature dogleg",
                [
                    "Formula: cos(Dogleg) = cos(Inc1) x cos(Inc2) + sin(Inc1) x sin(Inc2) x cos(Azi2 - Azi1)",
                    "Formula: RatioFactor = 1 when Dogleg ~= 0 else (2 / Dogleg) x tan(Dogleg / 2)",
                    "Source: UG_ST_navigation/.../survey_controller.dart _recalculateAllRows()",
                ],
            ),
            (
                "Survey TVD, North, East",
                [
                    "Formula: dTVD = (dMD / 2) x (cos(Inc1) + cos(Inc2)) x RatioFactor",
                    "Formula: dNorth = (dMD / 2) x (sin(Inc1)cos(Azi1) + sin(Inc2)cos(Azi2)) x RatioFactor",
                    "Formula: dEast = (dMD / 2) x (sin(Inc1)sin(Azi1) + sin(Inc2)sin(Azi2)) x RatioFactor",
                    "Formula: TVD, North, and East are cumulative sums of their interval deltas",
                    "Source: UG_ST_navigation/.../survey_controller.dart _recalculateAllRows()",
                ],
            ),
            (
                "Survey vertical section and DLS",
                [
                    "Formula: VSec = sqrt(North^2 + East^2) when ProjectAzimuth is disabled",
                    "Formula: VSec = North x cos(ProjectAzimuth) + East x sin(ProjectAzimuth) when enabled",
                    "Formula: DoglegSeverity = Dogleg_deg x IntervalBase / dMD, where IntervalBase is 100 ft or 30 m by unit setting",
                    "Source: UG_ST_navigation/.../survey_controller.dart _calculateVerticalSection()",
                ],
            ),
            (
                "Recap concentration",
                [
                    "Formula: Concentration = (Quantity x FactorPerPack) / SystemVolume",
                    "Formula: ton factor -> amount x 2000 lb; kg factor -> amount x 2.20462 lb; gal factor -> amount gal",
                    "Notes: Output unit depends on pack class, usually lb/bbl or gal/bbl.",
                    "Source: report/recap_tabs/concentration/controller/recap_concentration_controller.dart",
                ],
            ),
            (
                "Daily cost percentages",
                [
                    "Formula: PercentShare = (GroupAmount / TotalAmount) x 100",
                    "Source: daily_report/tabs/daily_cost/tabs/dailycost_percentagetable.dart",
                ],
            ),
        ],
    },
]


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        element = tc_mar.find(qn(f"w:{tag}"))
        if element is None:
            element = OxmlElement(f"w:{tag}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "DADCE0")


def set_paragraph_format(paragraph, after=8, before=0, line=1.15, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    fmt = paragraph.paragraph_format
    fmt.space_after = Pt(after)
    fmt.space_before = Pt(before)
    fmt.line_spacing = line
    paragraph.alignment = alignment


def set_run(run, *, size=11, bold=False, color="000000", name="Arial", italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_body_paragraph(document, text, *, after=8):
    paragraph = document.add_paragraph()
    set_paragraph_format(paragraph, after=after)
    run = paragraph.add_run(text)
    set_run(run, size=11)
    return paragraph


def add_heading(document, text, *, level=1):
    paragraph = document.add_paragraph()
    if level == 1:
        set_paragraph_format(paragraph, before=20, after=6)
        run = paragraph.add_run(text)
        set_run(run, size=20, bold=False)
    elif level == 2:
        set_paragraph_format(paragraph, before=18, after=6)
        run = paragraph.add_run(text)
        set_run(run, size=16, bold=False)
    else:
        set_paragraph_format(paragraph, before=16, after=4)
        run = paragraph.add_run(text)
        set_run(run, size=14, bold=False, color="434343")
    return paragraph


def add_title_block(document):
    title = document.add_paragraph()
    set_paragraph_format(title, after=3, before=0)
    title_run = title.add_run("MUDPRO+ Formula Reference")
    set_run(title_run, size=26, bold=False)

    meta = document.add_paragraph()
    set_paragraph_format(meta, after=8)
    run = meta.add_run(
        "Codebase snapshot reviewed on July 6, 2026. This document summarizes the formulas currently implemented across the desktop app and backend."
    )
    set_run(run, size=11, color="555555")

    add_body_paragraph(
        document,
        "Use this as a technical reference for report calculations, engineering helpers, pit balancing, survey math, hydraulics, and inventory-volume conversions.",
    )


def add_formula_table(document, rows):
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table)

    widths = [Inches(1.875), Inches(4.625)]
    header_cells = table.rows[0].cells
    for idx, width in enumerate(widths):
        header_cells[idx].width = width
        set_cell_margins(header_cells[idx])
        header_cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    left_header = header_cells[0].paragraphs[0]
    set_paragraph_format(left_header, after=4)
    left_run = left_header.add_run("Calculation")
    set_run(left_run, size=11, bold=True)

    right_header = header_cells[1].paragraphs[0]
    set_paragraph_format(right_header, after=4)
    right_run = right_header.add_run("Formula, Notes, And Source")
    set_run(right_run, size=11, bold=True)

    for title, details in rows:
        row_cells = table.add_row().cells
        for idx, width in enumerate(widths):
            row_cells[idx].width = width
            set_cell_margins(row_cells[idx])
            row_cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

        left_para = row_cells[0].paragraphs[0]
        set_paragraph_format(left_para, after=4)
        left_run = left_para.add_run(title)
        set_run(left_run, size=11, bold=True)

        first = True
        for detail in details:
            para = row_cells[1].paragraphs[0] if first else row_cells[1].add_paragraph()
            first = False
            set_paragraph_format(para, after=4)
            if ":" in detail:
                label, value = detail.split(":", 1)
                label_run = para.add_run(f"{label.strip()}: ")
                set_run(label_run, size=11, bold=True)
                value_run = para.add_run(value.strip())
                set_run(value_run, size=11)
            else:
                run = para.add_run(detail)
                set_run(run, size=11)

    return table


def configure_document(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)

    document.core_properties.title = "MUDPRO+ Formula Reference"
    document.core_properties.subject = "Formula and calculation reference for MudPro software"
    document.core_properties.language = "en-US"


def build_document():
    document = Document()
    configure_document(document)
    add_title_block(document)

    for section in SECTIONS:
        add_heading(document, section["title"], level=1)
        add_formula_table(document, section["rows"])

    add_heading(document, "9. Usage Notes", level=1)
    add_body_paragraph(
        document,
        "If a displayed value does not match expectations, verify the active report scope, carry-over source, unit system, and whether the UI is showing calculated volume, measured volume, or pending adjustment volume.",
    )
    add_body_paragraph(
        document,
        "Several recap, export, and pit-snapshot screens reuse the same backend formulas. When two screens disagree, the scope and source data usually differ before the formula does.",
    )

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
