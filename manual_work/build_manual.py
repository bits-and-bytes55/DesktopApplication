from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\DesktopApplication")
OUT_DIR = ROOT / "MudPro-Desktop-App" / "docs"
OUT_PATH = OUT_DIR / "MudPro_User_Manual_v1.0.docx"
SCREEN_DIR = Path(r"C:\Users\User\Pictures\Screenshots")

SCREENSHOTS = {
    "workspace": SCREEN_DIR / "Screenshot 2026-06-30 134514.png",
    "company_products": SCREEN_DIR / "Screenshot 2026-06-30 122523.png",
    "company_services": SCREEN_DIR / "Screenshot 2026-06-30 123307.png",
    "well_setup": SCREEN_DIR / "Screenshot 2026-06-30 140046.png",
    "report_well": SCREEN_DIR / "Screenshot 2026-07-02 094251.png",
    "report_pump": SCREEN_DIR / "Screenshot 2026-07-02 104315.png",
    "consume_product": SCREEN_DIR / "Screenshot 2026-07-03 132959.png",
    "return_product": SCREEN_DIR / "Screenshot 2026-07-03 140618.png",
    "return_lost": SCREEN_DIR / "Screenshot 2026-07-04 100031.png",
    "report_manager": SCREEN_DIR / "Screenshot 2026-07-02 154418.png",
}

BLUE = "6C9BCF"
DARK_BLUE = "1F4D78"
NAVY = "17365D"
LIGHT_BLUE = "EAF3FC"
TABLE_HEADER = "E8EEF5"
PALE_YELLOW = "FFF7CC"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
GRID = "CFE0F2"
BLACK = "000000"
WHITE = "FFFFFF"
GREEN = "2E7D32"
RED = "B42318"
GOLD = "7A5A00"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[min(idx, len(widths_dxa) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=11, color=BLACK, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, size=9, color=MID_GRAY)


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("MUDPRO+  |  User Manual")
    set_run_font(hr, size=9, color=MID_GRAY, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = fp.add_run("Page ")
    set_run_font(fr, size=9, color=MID_GRAY)
    add_page_field(fp)


def add_title(doc, text, size=30, color=NAVY, after=8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=after, line=1.0)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=True)
    return p


def add_subtitle(doc, text, size=14, color=DARK_BLUE, after=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=after, line=1.15)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color)
    return p


def add_body(doc, text, bold_prefix=None, color=BLACK, italic=False):
    p = doc.add_paragraph()
    set_paragraph_spacing(p)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True, color=color)
        r2 = p.add_run(text[len(bold_prefix) :])
        set_run_font(r2, color=color, italic=italic)
    else:
        run = p.add_run(text)
        set_run_font(run, color=color, italic=italic)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_step(doc, number, title, detail):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=4, after=3, line=1.2)
    set_keep_with_next(p)
    r1 = p.add_run(f"Step {number}: {title}")
    set_run_font(r1, size=11.5, color=DARK_BLUE, bold=True)
    d = doc.add_paragraph()
    set_paragraph_spacing(d, after=6)
    d.paragraph_format.left_indent = Inches(0.18)
    r2 = d.add_run(detail)
    set_run_font(r2)


def add_note(doc, label, text, fill=LIGHT_BLUE, accent=DARK_BLUE):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=5, after=7, line=1.18)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    border = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), accent)
    border.append(left)
    p_pr.append(border)
    r1 = p.add_run(f"{label}: ")
    set_run_font(r1, color=accent, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2)
    return p


def add_table(doc, headers, rows, widths_dxa, header_fill=TABLE_HEADER):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, text in enumerate(headers):
        set_cell_shading(header.cells[idx], header_fill)
        p = header.cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, after=0, line=1.0)
        run = p.add_run(text)
        set_run_font(run, size=10, bold=True)
    for row_index, row_data in enumerate(rows):
        cells = table.add_row().cells
        for idx, text in enumerate(row_data):
            if row_index % 2 == 1:
                set_cell_shading(cells[idx], "F9FBFD")
            p = cells[idx].paragraphs[0]
            p.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if idx == 0 and len(headers) > 2 else WD_ALIGN_PARAGRAPH.LEFT
            )
            set_paragraph_spacing(p, after=0, line=1.05)
            run = p.add_run(str(text))
            set_run_font(run, size=9.5)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def set_picture_alt_text(shape, text):
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", text)
    doc_pr.set("title", text)


def add_figure(doc, key, caption, width=6.3):
    path = SCREENSHOTS[key]
    if not path.exists():
        add_note(doc, "Screenshot unavailable", f"{caption} ka image file nahi mila.", fill="FFF1F0", accent=RED)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    shape = p.add_run().add_picture(str(path), width=Inches(width))
    set_picture_alt_text(shape, caption)
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(cp, after=8, line=1.0)
    run = cp.add_run(caption)
    set_run_font(run, size=9, color=MID_GRAY, italic=True)


def add_page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def build_manual():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    doc.core_properties.title = "MUDPRO+ User Manual"
    doc.core_properties.subject = "Desktop software operator guide"
    doc.core_properties.author = "MUDPRO+"
    doc.core_properties.keywords = "MUDPRO+, drilling, mud, daily report, user manual"

    # Cover: editorial_cover pattern with a real product screen.
    doc.add_paragraph().paragraph_format.space_after = Pt(44)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(kicker, after=14, line=1.0)
    kr = kicker.add_run("DESKTOP OPERATIONS GUIDE")
    set_run_font(kr, size=11, color=BLUE, bold=True)
    add_title(doc, "MUDPRO+ User Manual")
    add_subtitle(doc, "Pad setup se daily report generation tak step-by-step guide", after=8)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(meta, after=18, line=1.0)
    mr = meta.add_run("Version 1.0  |  July 2026  |  Windows Desktop")
    set_run_font(mr, size=10, color=MID_GRAY, bold=True)
    add_figure(
        doc,
        "workspace",
        "MUDPRO+ main workspace: Pad/Well/Report tree aur configuration tabs.",
        width=6.15,
    )
    add_note(
        doc,
        "Manual language",
        "Software labels English me hain; instructions simple Hinglish me diye gaye hain.",
    )

    add_page_break(doc)
    doc.add_heading("Contents", level=1)
    contents = [
        "1. Manual ka purpose aur workflow",
        "2. Interface aur common controls",
        "3. First-time Company Setup",
        "4. Pad, Well aur Report banana",
        "5. Report - Well tab",
        "6. Report - Mud tab",
        "7. Report - Pump tab",
        "8. Report - Operation tab",
        "9. Report - Pit tab aur balancing",
        "10. Report - Remarks tab",
        "11. Save, Lock, Calculate aur Report Generation",
        "12. Report Manager, Recap aur Utility",
        "13. Common issues aur troubleshooting",
        "14. Daily operator checklist",
    ]
    for item in contents:
        add_bullet(doc, item)
    add_note(
        doc,
        "Screen colors",
        "Blue section header hai, light-blue column header hai, pale-yellow editable/value area hai, aur gray/white read-only ya calculated area ho sakta hai. Lock state ke hisab se appearance change ho sakti hai.",
    )

    doc.add_heading("1. Manual ka purpose aur workflow", level=1)
    add_body(
        doc,
        "Ye manual mud engineer, drilling engineer, operator aur reporting team ko MUDPRO+ me master data setup, daily data entry, calculations, pit reconciliation aur report tools use karne me help karta hai.",
    )
    doc.add_heading("Recommended end-to-end sequence", level=2)
    for text in (
        "Mud Company Setup complete karein.",
        "Pad details enter karke save karein.",
        "Well banayein aur Well/Casing/Interval/Plan/Survey setup karein.",
        "New Report create karein.",
        "Report tabs ko sequence me complete karein: Well -> Mud -> Pump -> Operation -> Pit -> Remarks.",
        "Pit differences aur Storage measured values verify karein.",
        "Save, Calculate aur Lock karein; phir next report generate karein.",
    ):
        add_number(doc, text)
    add_note(
        doc,
        "Important",
        "Report aur operation calculations data-dependent hain. Empty mandatory setup, wrong units, ya unsaved edits calculation ko affect kar sakte hain.",
        fill="FFF8E8",
        accent=GOLD,
    )

    doc.add_heading("2. Interface aur common controls", level=1)
    add_figure(
        doc,
        "workspace",
        "Figure 1. Home workspace me selected Pad, Well, reports aur configuration tabs.",
    )
    doc.add_heading("Primary navigation", level=2)
    add_table(
        doc,
        ["Tab", "Use"],
        [
            ("Home", "Pad, Well, Inventory, Pit, Pump, SCE, Formation, Report aur Alert configuration."),
            ("Report", "Report Manager, Well Comparison, Recap aur Cost of Pad."),
            ("Utility", "Engineering Tools aur Unit Conversion."),
            ("Help", "User Manual, About, Disclaimer aur Abbreviation."),
        ],
        [1800, 7560],
    )
    doc.add_heading("Top toolbar", level=2)
    add_table(
        doc,
        ["Control", "Purpose"],
        [
            ("New Well", "Selected and completed Pad ke andar new Well banata hai."),
            ("Open Folder", "Saved/exported files location open karta hai."),
            ("Save", "Current active section ka latest data save karta hai."),
            ("Save As", "Current data ki separate copy/export flow ke liye."),
            ("Carry-over Pad", "Pad-level carry-over action."),
            ("New Report", "Selected Well ke liye next report create karta hai."),
            ("Carry-over", "Previous report data ko selected report me carry forward karta hai."),
            ("Lock / Unlock", "Report editing enable/disable karta hai."),
            ("Calculate", "Current tab ke calculation/save workflow ko execute karta hai."),
            ("Options", "Units aur application options."),
            ("Mud Company", "Company master-data setup open karta hai."),
            ("Upload / Batch Upload", "Supported data upload workflows."),
        ],
        [2100, 7260],
    )
    add_note(
        doc,
        "Selection rule",
        "Report, Utility aur kuch toolbar actions tabhi active hote hain jab valid Pad, Well aur Report selected ho.",
    )

    add_page_break(doc)
    doc.add_heading("3. First-time Company Setup", level=1)
    add_body(
        doc,
        "Home toolbar me Mud Company icon se Company Setup open karein. Engineers, Product, Services, Operator aur Others sections master data define karte hain. Ye data baad me dropdowns aur operation screens me use hota hai.",
    )
    add_figure(
        doc,
        "company_products",
        "Figure 2. Product master: Product, Code, SG, Unit, Group aur pricing fields.",
    )
    doc.add_heading("3.1 Products", level=2)
    add_step(doc, 1, "Editable mode open karein", "Top-right Editable control use karein.")
    add_step(doc, 2, "Product data enter karein", "Product name, unique Code, SG, unit quantity/class, Group aur price values fill karein.")
    add_step(doc, 3, "Save karein", "Bottom Save button ya available global save flow use karein.")
    add_step(doc, 4, "Verify karein", "Product list alphabetic order me visible honi chahiye aur report operations ke dropdown me available honi chahiye.")
    add_figure(
        doc,
        "company_services",
        "Figure 3. Package, Services aur Engineering master sections.",
    )
    doc.add_heading("3.2 Services and Engineering", level=2)
    add_bullet(doc, "Package me Name, Code, Unit aur Price define karein.")
    add_bullet(doc, "Services aur Engineering items ko unke relevant section me enter karein.")
    add_bullet(doc, "Existing row par right-click karke available edit/delete actions use karein.")
    add_bullet(doc, "Delete action successful hone par item database aur UI dono se remove hona chahiye.")
    doc.add_heading("3.3 Engineers, Operators and Others", level=2)
    add_bullet(doc, "Engineer/operator names aur contact details clean, unique format me rakhein.")
    add_bullet(doc, "Units, activities, mud types aur supporting lookup values ko Others me maintain karein.")
    add_bullet(doc, "Master data change ke baad affected report dropdown ko refresh/reopen karein.")

    add_page_break(doc)
    doc.add_heading("4. Pad, Well aur Report banana", level=1)
    doc.add_heading("4.1 Pad setup", level=2)
    add_figure(
        doc,
        "workspace",
        "Figure 4. Pad tab, left report tree aur top-level setup tabs.",
    )
    add_step(doc, 1, "Pad select/create karein", "Left tree me Pad select karein. Required details complete kiye bina Well creation enabled nahi hoga.")
    add_step(doc, 2, "Location choose karein", "Land ya Offshore select karein.")
    add_step(doc, 3, "Pad fields fill karein", "Field/Block, Rig, location, stock point, phone, operator, contractor aur representative fields enter karein.")
    add_step(doc, 4, "Logo and Memo", "Optional client logo aur memo add karein.")
    add_step(doc, 5, "Global Save", "Top toolbar Save se Pad data save karein.")
    add_note(
        doc,
        "Pad tabs",
        "PAD, INVENTORY, PIT, PUMP, SCE, FORMATION, REPORT aur ALERT configuration Pad/Well context ke liye use hote hain.",
    )

    doc.add_heading("4.2 Well setup", level=2)
    add_figure(
        doc,
        "well_setup",
        "Figure 5. Well setup tabs: Well, Casing, Interval, Plan aur Survey.",
    )
    add_step(doc, 1, "New Well click karein", "Completed Pad selected hone par toolbar ka New Well button active hota hai.")
    add_step(doc, 2, "Basic fields enter karein", "Well Name, API Well No. aur Spud Date fill karke Create karein.")
    add_step(doc, 3, "Well tab complete karein", "Location/reference values aur units verify karein.")
    add_step(doc, 4, "Casing define karein", "Description, Type, OD, Weight, ID, Top, Shoe aur Bit data enter karein.")
    add_step(doc, 5, "Interval/Plan/Survey complete karein", "Hole interval, mud plan, depth plan aur survey stations maintain karein.")
    add_note(
        doc,
        "Persistent setup",
        "Well/Casing/Pit configuration operational master setup hai. Data ko intentional delete action ke bina remove nahi karna chahiye.",
    )

    doc.add_heading("4.3 New Report", level=2)
    add_step(doc, 1, "Well select karein", "Left tree me active Well select karein.")
    add_step(doc, 2, "New Report click karein", "Toolbar me document icon se next report create karein.")
    add_step(doc, 3, "Report select karein", "Left tree me timestamp/report number wali new report row select karein.")
    add_step(doc, 4, "Carry-over decide karein", "Previous report ka setup/data chahiye to Carry-over use karein; confirmation message dhyan se padhein.")

    add_page_break(doc)
    doc.add_heading("5. Report - Well tab", level=1)
    add_figure(
        doc,
        "report_well",
        "Figure 6. Report Well tab ke main sections.",
    )
    add_body(
        doc,
        "Well tab daily drilling context define karta hai. Is tab ke values downstream hole volume, hydraulics, pit volume aur report summaries ko affect karte hain.",
    )
    add_table(
        doc,
        ["Section", "What to enter / verify"],
        [
            ("General", "Report no., date/time, engineers, representatives, activity, MD/TVD, inclination, drilling parameters aur temperatures."),
            ("Cased Hole", "Current casing description, OD, ID, top, shoe aur length."),
            ("Open Hole", "Hole description, ID, MD aur washout percentage."),
            ("Bit", "Manufacturer/type, bit count, size aur depth-in."),
            ("Nozzle", "Nozzle count aur size (1/32 in); TFA automatically calculate hota hai."),
            ("Drill String", "Description, OD, weight, ID, grade aur length. Total Length verify karein."),
            ("Time Distribution", "Daily activities aur minutes."),
        ],
        [1800, 7560],
    )
    add_note(
        doc,
        "Calculation dependency",
        "Hole volume ke liye casing/open-hole geometry, MD aur drill-string OD/ID/length complete hona zaroori hai.",
        fill="FFF8E8",
        accent=GOLD,
    )
    add_note(
        doc,
        "Nozzle",
        "Nozzle row update karne ke baad TFA change verify karein aur report generate karne se pehle Save karein.",
    )

    doc.add_heading("6. Report - Mud tab", level=1)
    add_body(
        doc,
        "Mud tab me Mud Properties, Rheology Model, rheology readings, Specific Gravity aur Solids enter kiye jate hain.",
    )
    add_step(doc, 1, "Mud Properties", "Sample details, temperatures, depth, MW, funnel viscosity aur relevant chemical/physical properties enter karein.")
    add_step(doc, 2, "Rheology Model select karein", "Available model, jaise Bingham, select karein.")
    add_step(doc, 3, "Rheology readings", "600/300/200/100/6/3 RPM readings aur calculated PV/YP verify karein.")
    add_step(doc, 4, "Specific Gravity and Solids", "Oil/HGS/LGS aur shale/bentonite CEC values enter karein.")
    add_step(doc, 5, "Save and review", "Units aur decimal format verify karke Save karein.")
    add_note(
        doc,
        "Unit caution",
        "Mud weight, viscosity aur concentration units Options me configured units se match hone chahiye.",
        fill="FFF8E8",
        accent=GOLD,
    )

    add_page_break(doc)
    doc.add_heading("7. Report - Pump tab", level=1)
    add_figure(
        doc,
        "report_pump",
        "Figure 7. Pump, Summary, Shaker aur Other SCE sections.",
    )
    add_table(
        doc,
        ["Section", "Use"],
        [
            ("Pump", "Model/type, liner ID, rod OD, stroke length, efficiency, displacement, stroke rate aur flow rate."),
            ("Summary", "Pump Rate, Pump Pressure aur Downhole Tools Pressure Loss."),
            ("Shaker", "Shaker/model, screen values, operating time aur OOC weight."),
            ("Other SCE", "Other solids-control equipment, model, U/F, O/F, time aur OOC weight."),
        ],
        [1800, 7560],
    )
    add_step(doc, 1, "Configured equipment choose karein", "Pad PUMP/SCE setup me available equipment dropdown se select karein.")
    add_step(doc, 2, "Operating inputs enter karein", "Stroke, screen, flow aur time values fill karein.")
    add_step(doc, 3, "Calculated outputs verify karein", "Displacement/rate/summary values expected range me hone chahiye.")
    add_step(doc, 4, "Save karein", "Pump tab leave karne ya report generate karne se pehle save complete hone dein.")

    doc.add_heading("8. Report - Operation tab", level=1)
    add_body(
        doc,
        "Operation tab daily material aur mud movements record karta hai. Left operation menu me row select karke operation choose karein. Har operation apna independent data rakhta hai; unrelated operation values merge nahi honi chahiye.",
    )
    add_table(
        doc,
        ["Operation", "Purpose"],
        [
            ("Consume Services", "Package, service aur engineering usage record."),
            ("Consume Product", "Inventory product usage, cost aur volume addition."),
            ("Receive Product", "Product/package inventory receive karna."),
            ("Return Product", "Product/package inventory return karna."),
            ("Transfer Mud", "Pits/systems ke beech mud transfer."),
            ("Receive Mud", "BOL ke against premixed mud receive karna."),
            ("Return / Lost Mud", "Returned ya lost mud transaction."),
            ("Add Water", "Selected system me water volume add karna."),
            ("Switch Pit", "Active/storage pit status switch."),
            ("Switch Mud Type", "Active pits/storage displacement workflow."),
            ("Empty Active System", "Active system dump ya storage transfer."),
            ("Other Vol. Addition - Active System", "Other volume addition to active system."),
            ("Mud Loss - Active System", "Active system mud loss."),
            ("Mud Loss - Storage", "Storage mud loss."),
        ],
        [3000, 6360],
    )

    doc.add_heading("8.1 Consume Product", level=2)
    add_figure(
        doc,
        "consume_product",
        "Figure 8. Consume Product table, distribution aur Add Water controls.",
    )
    add_step(doc, 1, "Select Products", "Inventory list se one or more products choose karein.")
    add_step(doc, 2, "Input Method", "Used ya Final method select karein.")
    add_step(doc, 3, "Usage enter karein", "Initial, Adjust aur Used/Final values fill karein; Cost aur Vol (bbl) verify karein.")
    add_step(doc, 4, "Distribution", "Active System ya selected pit ko calculated volume distribute karein.")
    add_step(doc, 5, "Save", "Operation tab ke save/calculate flow se product rows aur distribution save karein.")
    add_note(
        doc,
        "Carry forward",
        "Previous report ka Final, same product select karne par next report ka Initial banta hai.",
    )

    doc.add_heading("8.2 Receive and Return Product", level=2)
    add_figure(
        doc,
        "return_product",
        "Figure 9. Return Product me BOL, Product aur Package sections.",
    )
    add_bullet(doc, "BOL number enter karein.")
    add_bullet(doc, "Product/Package select karke quantity/unit details complete karein.")
    add_bullet(doc, "Return All Inventory sirf intentional full-return case me use karein.")
    add_bullet(doc, "Save ke baad inventory balance verify karein.")

    doc.add_heading("8.3 Return / Lost Mud", level=2)
    add_figure(
        doc,
        "return_lost",
        "Figure 10. Return/Lost Mud fields.",
    )
    add_bullet(doc, "Premixed Mud applicable ho to checkbox aur source select karein.")
    add_bullet(doc, "From, To, returned/lost volume, MW, Mud Type, BOL aur cost values enter karein.")
    add_bullet(doc, "Leased status applicable ho to checkbox use karein.")

    add_page_break(doc)
    doc.add_heading("9. Report - Pit tab aur balancing", level=1)
    add_body(
        doc,
        "Pit tab physical pit measurements ko calculated report volume se reconcile karta hai. Is tab ko report closeout se pehle zaroor verify karein.",
    )
    add_table(
        doc,
        ["Section", "Use"],
        [
            ("Active Pits", "Pit-wise measured volume, MW aur mud selection."),
            ("Storage", "Calculated volume aur measured volume comparison."),
            ("Volume Name", "Hole, Active Pits, Active System, End Vol. aur pending difference."),
            ("Haul Off", "Loads, volume, weight, oil/water/solids aur OOC weight."),
            ("Pit Snapshot", "Hole volume, pit distribution aur concentration snapshot."),
        ],
        [1800, 7560],
    )
    doc.add_heading("9.1 Core relationships", level=2)
    add_note(doc, "Active System", "Hole + Active Pits = Active System.")
    add_note(doc, "Pending difference", "End Vol. - Active System batata hai ki pit adjustment abhi pending hai.")
    add_note(doc, "Target", "Pit adjustment ke baad End Vol. aur Active System equal hone chahiye; pending difference 0 hona chahiye.")
    doc.add_heading("9.2 Positive and negative adjustment", level=2)
    add_bullet(doc, "Positive pending: required amount kisi active pit me add karein.")
    add_bullet(doc, "Negative pending: required amount kisi active pit se reduce karein.")
    add_bullet(doc, "Pit capacity se zyada measured volume enter nahi karein.")
    add_bullet(doc, "Adjustment current report me hona chahiye; previous report values ko unintentionally modify nahi karein.")
    doc.add_heading("9.3 Storage validation", level=2)
    add_body(
        doc,
        "Storage me Calculated Vol aur Measured Vol equal hone par report normal generate hoti hai. Missing/different value par confirmation popup aata hai. Yes se report continue hoti hai; No se generation cancel hoti hai.",
    )
    doc.add_heading("9.4 Pit Snapshot", level=2)
    add_bullet(doc, "Hole Volume (bbl) me String, Annulus, Below bit, Hole aur Displacement review karein.")
    add_bullet(doc, "Displacement drill-string geometry ke OD volume aur ID/internal volume difference se calculate hota hai.")
    add_bullet(doc, "Active Pits/Storage distribution aur product concentration values verify karein.")

    doc.add_heading("10. Report - Remarks tab", level=1)
    add_table(
        doc,
        ["Field", "Use"],
        [
            ("Recommended Tour Treatments", "Next tour/team ke liye recommended treatment."),
            ("Operational Comments", "Current operation ki important observations."),
            ("Remarks", "Report recap/general remarks."),
            ("Internal Notes", "Internal-only notes; external distribution se pehle review karein."),
        ],
        [2600, 6760],
    )
    add_step(doc, 1, "Relevant text enter karein", "Clear, factual aur time-specific statements use karein.")
    add_step(doc, 2, "Attachment", "Supported attachment control available ho to evidence/document attach karein.")
    add_step(doc, 3, "Save", "Tab leave karne ya report lock karne se pehle Save karein.")

    add_page_break(doc)
    doc.add_heading("11. Save, Lock, Calculate aur Report Generation", level=1)
    doc.add_heading("11.1 Save", level=2)
    add_bullet(doc, "Top global Save current active module ke latest edits persist karta hai.")
    add_bullet(doc, "Cell edit ke turant baad tab switch/report generation na karein jab saving indicator active ho.")
    add_bullet(doc, "Master setup aur report data alag scopes me save hote hain.")
    doc.add_heading("11.2 Lock / Unlock", level=2)
    add_bullet(doc, "Unlock state me editable fields active hote hain.")
    add_bullet(doc, "Lock state accidental edits prevent karta hai.")
    add_bullet(doc, "Final review ke baad report lock karein.")
    doc.add_heading("11.3 Calculate", level=2)
    add_bullet(doc, "Current tab calculations refresh karta hai.")
    add_bullet(doc, "Calculated values ko source inputs aur units ke against verify karein.")
    doc.add_heading("11.4 New Report generation checklist", level=2)
    for text in (
        "Well geometry and MD/TVD complete.",
        "Mud and rheology values saved.",
        "Pump/SCE operating data saved.",
        "All selected operations saved.",
        "Pit pending difference zero or intentionally accepted.",
        "Storage calculated/measured volumes checked.",
        "Remarks completed.",
        "Current report saved and locked.",
    ):
        add_bullet(doc, text)
    add_note(
        doc,
        "Do not proceed",
        "Red error panel, overflow warning, missing mandatory selection ya failed save message visible ho to report generation se pehle issue resolve karein.",
        fill="FFF1F0",
        accent=RED,
    )

    doc.add_heading("12. Report Manager, Recap aur Utility", level=1)
    add_figure(
        doc,
        "report_manager",
        "Figure 11. Report Manager search criteria aur result area.",
    )
    doc.add_heading("12.1 Report Manager", level=2)
    add_step(doc, 1, "Current Well choose karein", "Dropdown me required Well select karein.")
    add_step(doc, 2, "Criteria enable karein", "Date, Report No., Depth, MW ya text filters ke checkbox select karein.")
    add_step(doc, 3, "Search karein", "Search button se matching reports load karein.")
    add_step(doc, 4, "Result action", "Required report select/open karein; delete sirf authorized case me.")
    doc.add_heading("12.2 Other Report tools", level=2)
    add_bullet(doc, "Well Comparison: selected wells/reports ka comparison.")
    add_bullet(doc, "Recap: report data ka summarized analytical view.")
    add_bullet(doc, "Cost of Pad: Pad-level cost view.")
    doc.add_heading("12.3 Utility", level=2)
    add_bullet(doc, "Engineering Tools: hydraulics/rheology related calculations.")
    add_bullet(doc, "Unit Conversion: source value/unit se target unit conversion.")
    add_note(
        doc,
        "Engineering review",
        "Software calculation engineering judgment ka replacement nahi hai. Input quality, selected model aur unit system verify karna user ki responsibility hai.",
        fill="FFF8E8",
        accent=GOLD,
    )

    add_page_break(doc)
    doc.add_heading("13. Common issues aur troubleshooting", level=1)
    add_table(
        doc,
        ["Issue", "Check / action"],
        [
            ("Report tabs disabled", "Pad, Well aur Report selection complete karein."),
            ("Dropdown empty", "Company Setup/Inventory master data save hua hai ya nahi check karein; refresh/reopen karein."),
            ("Old value reload ho rahi hai", "Current report selected hai ya nahi verify karein; field edit ke baad Save complete hone dein."),
            ("Calculation zero", "Required geometry, unit, depth, volume aur selection inputs missing ho sakte hain."),
            ("Pit difference non-zero", "Active Pits ko positive/negative pending amount ke according adjust karein."),
            ("Measured Vol popup", "Storage calculated vs measured volumes compare karein; intentional mismatch ho tabhi Yes choose karein."),
            ("Data load nahi ho raha", "Network, backend/API availability aur selected Well/Report context check karein."),
            ("Red error screen", "Current action stop karein, screenshot/error text capture karein aur support ko report karein."),
            ("Text/controls clipped", "Window maximize karein; supported display scaling/resolution use karein."),
        ],
        [2700, 6660],
    )
    doc.add_heading("Support ko kya dena hai", level=2)
    add_bullet(doc, "Pad name, Well name aur Report number.")
    add_bullet(doc, "Exact tab/section/operation name.")
    add_bullet(doc, "Expected result aur actual result.")
    add_bullet(doc, "Full screenshot including error message.")
    add_bullet(doc, "Issue reproduce karne ke exact steps.")
    add_bullet(doc, "Backend/server log ka relevant timestamp, agar available ho.")

    doc.add_heading("14. Daily operator checklist", level=1)
    checklist = [
        ("Start of shift", "Correct Pad, Well aur latest Report selected."),
        ("Well", "MD/TVD, activity, hole/casing, bit/nozzle, drill string and time distribution updated."),
        ("Mud", "Properties, rheology, SG and solids updated."),
        ("Pump", "Pump, shaker and SCE inputs updated."),
        ("Operations", "All inventory/mud movements recorded and saved."),
        ("Pit", "Measured volumes updated; pending difference reconciled."),
        ("Storage", "Calculated and measured volume reviewed."),
        ("Remarks", "Operational comments, remarks and internal notes completed."),
        ("Closeout", "Save -> Calculate -> Review -> Lock."),
        ("Next report", "New Report/carry-over action only after current report closeout."),
    ]
    add_table(doc, ["Stage", "Verification"], checklist, [2200, 7160])
    add_note(
        doc,
        "Final rule",
        "Report generate karne se pehle data completeness, units, pit balance aur save status ko ek baar visually verify karein.",
        fill="EAF6EA",
        accent=GREEN,
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(18)
    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(end, before=12, after=0, line=1.0)
    er = end.add_run("End of MUDPRO+ User Manual")
    set_run_font(er, size=12, color=DARK_BLUE, bold=True)

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build_manual()
