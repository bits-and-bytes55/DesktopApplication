from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches


DOCX = Path(
    r"D:\DesktopApplication\MudPro-Desktop-App\docs\MudPro_User_Manual_v1.0.docx"
)


def int_attr(element, name):
    value = element.get(qn(name))
    return int(value) if value else 0


with ZipFile(DOCX) as archive:
    bad_member = archive.testzip()
    if bad_member:
        raise RuntimeError(f"Corrupt DOCX member: {bad_member}")
    media = [name for name in archive.namelist() if name.startswith("word/media/")]

doc = Document(DOCX)
headings = [
    paragraph.text.strip()
    for paragraph in doc.paragraphs
    if paragraph.style.name.startswith("Heading")
]
if not headings or any(not heading for heading in headings):
    raise RuntimeError("Heading structure is missing or contains empty headings")

full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
for marker in ("TODO", "PLACEHOLDER", "[[", "]]"):
    if marker in full_text:
        raise RuntimeError(f"Unresolved marker found: {marker}")

section = doc.sections[0]
expected = {
    "page_width": int(Inches(8.5)),
    "page_height": int(Inches(11)),
    "top_margin": int(Inches(1)),
    "right_margin": int(Inches(1)),
    "bottom_margin": int(Inches(1)),
    "left_margin": int(Inches(1)),
}
for name, value in expected.items():
    actual = int(getattr(section, name))
    if abs(actual - value) > 2:
        raise RuntimeError(f"Unexpected {name}: {actual}")

for table_index, table in enumerate(doc.tables, start=1):
    grid_widths = [
        int_attr(col, "w:w") for col in table._tbl.tblGrid.gridCol_lst
    ]
    if not grid_widths or sum(grid_widths) != 9360:
        raise RuntimeError(
            f"Table {table_index} grid width is {sum(grid_widths)}, expected 9360"
        )
    for row_index, row in enumerate(table.rows, start=1):
        cell_widths = []
        for cell in row.cells:
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            cell_widths.append(int_attr(tc_w, "w:w") if tc_w is not None else 0)
        if cell_widths != grid_widths:
            raise RuntimeError(
                f"Table {table_index} row {row_index} widths {cell_widths} "
                f"do not match grid {grid_widths}"
            )

print(
    {
        "file": str(DOCX),
        "size_bytes": DOCX.stat().st_size,
        "paragraphs": len(doc.paragraphs),
        "headings": len(headings),
        "tables": len(doc.tables),
        "embedded_images": len(media),
        "status": "structural QA passed",
    }
)
